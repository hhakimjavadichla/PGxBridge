"""
Word document report generator for PGx patient-facing reports.

Generates Word documents from extracted PGx data using a template.
"""

import os
import io
import copy
import logging
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)

# Path to template files
PATIENT_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "templates", "PGx Patient Report_04.08.25.docx")
EHR_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "templates", "Note template.docx")
EHR_TEMPLATE_JINJA2_PATH = os.path.join(os.path.dirname(__file__), "..", "templates", "Note_template_jinja2.docx")
PATIENT_TEMPLATE_JINJA2_PATH = os.path.join(os.path.dirname(__file__), "..", "templates", "PGx_Patient_Report_jinja2.docx")

# Legacy alias
TEMPLATE_PATH = PATIENT_TEMPLATE_PATH

# Medication examples by gene (from template)
MEDICATION_EXAMPLES = {
    "CYP2B6": "Infectious Diseases: efavirenz",
    "CYP2C9": "Cardiology: fluvastatin (Lescol), warfarin (Coumadin); Pain Management: celecoxib (Celebrex)",
    "CYP2C19": "Cardiology: clopidogrel (Plavix); Gastroenterology: dexlansoprazole (Dexilant), esomeprazole (Nexium), lansoprazole (Prevacid), omeprazole (Prilosec), pantoprazole (Protonix)",
    "CYP2C cluster": "Cardiology: warfarin (Coumadin)",
    "CYP2D6": "Cardiology: flecainide (Tambocor), metoprolol (Lopressor); Pain Management: codeine, tramadol (Ultram)",
    "CYP3A5": "Immunosuppression: tacrolimus (Prograf)",
    "CYP4F2": "Cardiology: warfarin (Coumadin)",
    "DPYD": "Oncology: fluorouracil, capecitabine (Xeloda); Infectious Diseases: flucytosine (Ancobon)",
    "HLA-A": "Neurology: carbamazepine (Tegretol), oxcarbazepine (Trileptal)",
    "HLA-B": "Neurology: carbamazepine (Tegretol); Infectious Diseases: abacavir (Ziagen); Rheumatology: allopurinol (Zyloprim)",
    "IFNL4": "Infectious Diseases: peginterferon alfa-2a (Pegasys)",
    "NUDT15": "Immunosuppression: azathioprine (Imuran), mercaptopurine (Purinethol)",
    "SLCO1B1": "Cardiology: atorvastatin (Lipitor), simvastatin (Zocor), rosuvastatin (Crestor)",
    "TPMT": "Immunosuppression: azathioprine (Imuran), mercaptopurine (Purinethol)",
    "UGT1A1": "Oncology: irinotecan (Camptosar); Infectious Diseases: atazanavir (Reyataz)",
    "VKORC1": "Cardiology: warfarin (Coumadin)",
    "NAT2": "Infectious Diseases: isoniazid",
}


def get_medication_examples(gene: str) -> str:
    """Get medication examples for a gene."""
    # Normalize gene name for lookup
    gene_upper = gene.upper()
    for key, value in MEDICATION_EXAMPLES.items():
        if key.upper() == gene_upper:
            return value
    return ""


def is_high_risk_or_cyp2c19(gene_data: Dict) -> bool:
    """
    Check if a gene result should be in Table 1.
    
    Criteria:
    - CPIC EHR Priority is "Abnormal/Priority/High Risk"
    - OR gene is CYP2C19 (all results regardless of priority)
    """
    gene = gene_data.get("gene", "").upper()
    ehr_priority = gene_data.get("cpic_ehr_priority", "") or ""
    
    # All CYP2C19 results go in Table 1
    if "CYP2C19" in gene:
        return True
    
    # High-risk variants go in Table 1
    if "high risk" in ehr_priority.lower() or "abnormal" in ehr_priority.lower():
        return True
    
    return False


def get_priority_category(gene_data: Dict) -> str:
    """
    Determine which table a gene belongs to.
    
    Returns:
        'priority' - Table 1 (high-risk + all CYP2C19)
        'standard' - Table 2 (normal/routine/low risk)
        'unknown' - Table 3 (unknown/indeterminate priority)
    """
    gene = gene_data.get("gene", "").upper()
    ehr_priority = (gene_data.get("cpic_ehr_priority", "") or "").lower()
    
    # All CYP2C19 results go to Priority (Table 1)
    if "CYP2C19" in gene:
        return "priority"
    
    # Check for high-risk indicators (must check before "normal" due to "abnormal" containing "normal")
    if "abnormal" in ehr_priority or "high risk" in ehr_priority or "priority" in ehr_priority:
        return "priority"
    
    # Check for normal/routine/low risk - use exact phrase matching to avoid "abnormal" issue
    # The standard CPIC format is "Normal/Routine/Low Risk"
    if ehr_priority.startswith("normal") or "routine" in ehr_priority or "low risk" in ehr_priority:
        return "standard"
    
    # Everything else is unknown (empty, "none", indeterminate, etc.)
    return "unknown"


def is_normal_risk(gene_data: Dict) -> bool:
    """
    Check if a gene result has Normal/Routine/Low Risk priority.
    Excludes CYP2C19 (which always goes to high-risk table).
    """
    return get_priority_category(gene_data) == "standard"


def is_unknown_priority(gene_data: Dict) -> bool:
    """
    Check if a gene result has unknown/indeterminate CPIC EHR priority.
    """
    return get_priority_category(gene_data) == "unknown"


def clear_table_keep_header(table):
    """Remove all rows from a table except the header row."""
    # Keep only the first row (header)
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)


def set_cell_border(cell, border_color="000000", border_size="4"):
    """
    Apply borders to a table cell.
    
    Args:
        cell: The table cell to apply borders to
        border_color: Hex color code for border (default black)
        border_size: Border width in eighths of a point (default 4 = 0.5pt)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Remove existing borders if any
    existing_borders = tcPr.find(qn('w:tcBorders'))
    if existing_borders is not None:
        tcPr.remove(existing_borders)
    
    # Create new borders element
    tcBorders = OxmlElement('w:tcBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:color'), border_color)
        border.set(qn('w:space'), '0')
        tcBorders.append(border)
    
    tcPr.append(tcBorders)


def add_row_to_table(table, cells_data: List[str], apply_borders: bool = True):
    """
    Add a new row to the table with the given cell data.
    
    Args:
        table: The table to add a row to
        cells_data: List of strings for each cell
        apply_borders: Whether to apply borders to cells (default True)
    """
    row = table.add_row()
    for i, text in enumerate(cells_data):
        if i < len(row.cells):
            row.cells[i].text = text
            if apply_borders:
                set_cell_border(row.cells[i])


def generate_patient_report(
    patient_info: Dict,
    pgx_genes: List[Dict],
    output_path: Optional[str] = None
) -> bytes:
    """
    Generate a patient-facing PGx report Word document.
    
    Args:
        patient_info: Dictionary with patient information
        pgx_genes: List of gene data dictionaries with CPIC annotations
        output_path: Optional path to save the document (for debugging)
        
    Returns:
        bytes: The generated Word document as bytes
    """
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found at: {TEMPLATE_PATH}")
    
    # Load template
    doc = Document(TEMPLATE_PATH)
    
    # === POPULATE PATIENT INFO (Table 0) ===
    patient_table = doc.tables[0]
    
    # Row 0: Patient name and Laboratory
    patient_name = patient_info.get("patient_name", "")
    if patient_name:
        patient_table.rows[0].cells[1].text = patient_name
    
    # Row 1: Date of birth and Date of test
    dob = patient_info.get("date_of_birth", "")
    if dob:
        patient_table.rows[1].cells[1].text = dob
    
    report_date = patient_info.get("report_date", "")
    if report_date:
        patient_table.rows[1].cells[3].text = report_date
    
    # Row 2: Date of interpretation (use report_date if available)
    if report_date:
        patient_table.rows[2].cells[3].text = report_date
    
    # === POPULATE TABLE 1: HIGH-RISK + CYP2C19 RESULTS (Table 3) ===
    table1 = doc.tables[3]
    
    # Clear existing data rows, keep header
    clear_table_keep_header(table1)
    
    # Filter genes for Table 1
    table1_genes = [g for g in pgx_genes if is_high_risk_or_cyp2c19(g)]
    
    # Sort by gene name for consistent ordering
    table1_genes.sort(key=lambda x: x.get("gene", ""))
    
    # Add rows for each gene
    for gene_data in table1_genes:
        gene = gene_data.get("gene", "")
        genotype = gene_data.get("genotype", "")
        phenotype = gene_data.get("metabolizer_status", "") or gene_data.get("cpic_phenotype", "")
        medications = get_medication_examples(gene)
        
        add_row_to_table(table1, [gene, genotype, phenotype, medications])
    
    # === POPULATE TABLE 2: ALL PGX RESULTS (Table 4) ===
    table2 = doc.tables[4]
    
    # Clear existing data rows, keep header
    clear_table_keep_header(table2)
    
    # Sort all genes by gene name
    all_genes = sorted(pgx_genes, key=lambda x: x.get("gene", ""))
    
    # Add rows for each gene
    for gene_data in all_genes:
        gene = gene_data.get("gene", "")
        genotype = gene_data.get("genotype", "")
        phenotype = gene_data.get("metabolizer_status", "") or gene_data.get("cpic_phenotype", "")
        
        add_row_to_table(table2, [gene, genotype, phenotype])
    
    # === SAVE DOCUMENT ===
    if output_path:
        doc.save(output_path)
        logger.info(f"Report saved to: {output_path}")
    
    # Return as bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_patient_report_filename(patient_info: Dict) -> str:
    """Generate a filename for the patient report."""
    patient_name = patient_info.get("patient_name", "Unknown")
    report_date = patient_info.get("report_date", "")
    
    # Clean up patient name for filename
    safe_name = "".join(c if c.isalnum() or c in " -_" else "" for c in patient_name)
    safe_name = safe_name.replace(" ", "_")
    
    if report_date:
        safe_date = report_date.replace("/", "-").replace(" ", "")
        return f"PGx_Report_{safe_name}_{safe_date}.docx"
    else:
        return f"PGx_Report_{safe_name}.docx"


def generate_patient_report_docxtpl(
    patient_info: Dict,
    pgx_genes: List[Dict],
    output_path: Optional[str] = None
) -> bytes:
    """
    Generate a patient-facing PGx report using docxtpl (Jinja2 templating).
    
    This function uses a Jinja2-enabled Word template for better formatting
    preservation including table borders, fonts, and styles.
    
    Args:
        patient_info: Dictionary with patient information
        pgx_genes: List of gene data dictionaries with CPIC annotations
        output_path: Optional path to save the document (for debugging)
        
    Returns:
        bytes: The generated Word document as bytes
    """
    if not os.path.exists(PATIENT_TEMPLATE_JINJA2_PATH):
        raise FileNotFoundError(f"Jinja2 Patient template not found at: {PATIENT_TEMPLATE_JINJA2_PATH}")
    
    # Load template
    template = DocxTemplate(PATIENT_TEMPLATE_JINJA2_PATH)
    
    # Filter high-risk genes (high-risk + CYP2C19)
    high_risk_genes = []
    all_genes = []
    
    for g in pgx_genes:
        gene_row = {
            "gene": g.get("gene", ""),
            "genotype": g.get("genotype", ""),
            "phenotype": g.get("metabolizer_status", "") or g.get("cpic_phenotype", ""),
            "medications": get_medication_examples(g.get("gene", ""))
        }
        
        all_genes.append(gene_row)
        
        if is_high_risk_or_cyp2c19(g):
            high_risk_genes.append(gene_row)
    
    # Sort both lists by gene name
    high_risk_genes.sort(key=lambda x: x["gene"])
    all_genes.sort(key=lambda x: x["gene"])
    
    logger.info(f"Patient Report (docxtpl): {len(high_risk_genes)} high-risk, {len(all_genes)} total genes")
    
    # Build context for template rendering
    context = {
        # Patient info
        "patient_name": patient_info.get("patient_name", ""),
        "date_of_birth": patient_info.get("date_of_birth", ""),
        "report_date": patient_info.get("report_date", ""),
        
        # Gene lists
        "high_risk_genes": high_risk_genes,
        "all_genes": all_genes,
    }
    
    # Render the template with context
    template.render(context)
    
    # Save to file if path provided (for debugging)
    if output_path:
        template.save(output_path)
        logger.info(f"Patient report (docxtpl) saved to: {output_path}")
    
    # Return as bytes
    buffer = io.BytesIO()
    template.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_ehr_report(
    patient_info: Dict,
    pgx_genes: List[Dict],
    output_path: Optional[str] = None
) -> bytes:
    """
    Generate an EHR-facing PGx report Word document.
    
    Table structure:
    - Table 0: Patient info (Name, DOB, MRN, Insurance)
    - Table 1: Priority results (high-risk + all CYP2C19)
    - Table 2: Standard results (normal/routine/low risk)
    - Table 3: Unknown results (unknown/indeterminate priority)
    
    Args:
        patient_info: Dictionary with patient information
        pgx_genes: List of gene data dictionaries with CPIC annotations
        output_path: Optional path to save the document (for debugging)
        
    Returns:
        bytes: The generated Word document as bytes
    """
    if not os.path.exists(EHR_TEMPLATE_PATH):
        raise FileNotFoundError(f"EHR template not found at: {EHR_TEMPLATE_PATH}")
    
    # Load template
    doc = Document(EHR_TEMPLATE_PATH)
    
    # === POPULATE PATIENT INFO (Table 0) ===
    patient_table = doc.tables[0]
    
    # Row 0: Name, DOB, MRN, Insurance
    patient_name = patient_info.get("patient_name", "")
    dob = patient_info.get("date_of_birth", "")
    
    if patient_name:
        patient_table.rows[0].cells[0].text = f"Name: {patient_name}"
    if dob:
        patient_table.rows[0].cells[1].text = f"DOB: {dob}"
    
    # === CATEGORIZE GENES (mutually exclusive) ===
    priority_genes = [g for g in pgx_genes if get_priority_category(g) == "priority"]
    standard_genes = [g for g in pgx_genes if get_priority_category(g) == "standard"]
    unknown_genes = [g for g in pgx_genes if get_priority_category(g) == "unknown"]
    
    # Sort each list by gene name
    priority_genes.sort(key=lambda x: x.get("gene", ""))
    standard_genes.sort(key=lambda x: x.get("gene", ""))
    unknown_genes.sort(key=lambda x: x.get("gene", ""))
    
    logger.info(f"EHR Report: {len(priority_genes)} priority, {len(standard_genes)} standard, {len(unknown_genes)} unknown")
    
    # === POPULATE TABLE 1: PRIORITY RESULTS (Table 1) ===
    # 4 columns: Gene, Genotype, Phenotype/Clinical Implication, Examples of affected medications
    table1 = doc.tables[1]
    clear_table_keep_header(table1)
    
    for gene_data in priority_genes:
        gene = gene_data.get("gene", "")
        genotype = gene_data.get("genotype", "")
        phenotype = gene_data.get("metabolizer_status", "") or gene_data.get("cpic_phenotype", "")
        medications = get_medication_examples(gene)
        
        add_row_to_table(table1, [gene, genotype, phenotype, medications])
    
    # === POPULATE TABLE 2: STANDARD RESULTS (Table 2) ===
    # 3 columns: Gene, Genotype, Phenotype/Clinical Implication
    table2 = doc.tables[2]
    clear_table_keep_header(table2)
    
    for gene_data in standard_genes:
        gene = gene_data.get("gene", "")
        genotype = gene_data.get("genotype", "")
        phenotype = gene_data.get("metabolizer_status", "") or gene_data.get("cpic_phenotype", "")
        
        add_row_to_table(table2, [gene, genotype, phenotype])
    
    # === POPULATE TABLE 3: UNKNOWN RESULTS (Table 3) ===
    # 3 columns: Gene, Genotype, Phenotype/Clinical Implication
    table3 = doc.tables[3]
    clear_table_keep_header(table3)
    
    for gene_data in unknown_genes:
        gene = gene_data.get("gene", "")
        genotype = gene_data.get("genotype", "")
        phenotype = gene_data.get("metabolizer_status", "") or gene_data.get("cpic_phenotype", "")
        
        add_row_to_table(table3, [gene, genotype, phenotype])
    
    # === SAVE DOCUMENT ===
    if output_path:
        doc.save(output_path)
        logger.info(f"EHR report saved to: {output_path}")
    
    # Return as bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_ehr_report_filename(patient_info: Dict) -> str:
    """Generate a filename for the EHR report."""
    patient_name = patient_info.get("patient_name", "Unknown")
    report_date = patient_info.get("report_date", "")
    
    # Clean up patient name for filename
    safe_name = "".join(c if c.isalnum() or c in " -_" else "" for c in patient_name)
    safe_name = safe_name.replace(" ", "_")
    
    if report_date:
        safe_date = report_date.replace("/", "-").replace(" ", "")
        return f"PGx_EHR_Note_{safe_name}_{safe_date}.docx"
    else:
        return f"PGx_EHR_Note_{safe_name}.docx"


def generate_ehr_report_docxtpl(
    patient_info: Dict,
    pgx_genes: List[Dict],
    output_path: Optional[str] = None
) -> bytes:
    """
    Generate an EHR-facing PGx report using docxtpl (Jinja2 templating).
    
    This function uses a Jinja2-enabled Word template for better formatting
    preservation including table borders, fonts, and styles.
    
    Args:
        patient_info: Dictionary with patient information
        pgx_genes: List of gene data dictionaries with CPIC annotations
        output_path: Optional path to save the document (for debugging)
        
    Returns:
        bytes: The generated Word document as bytes
    """
    if not os.path.exists(EHR_TEMPLATE_JINJA2_PATH):
        raise FileNotFoundError(f"Jinja2 EHR template not found at: {EHR_TEMPLATE_JINJA2_PATH}")
    
    # Load template
    template = DocxTemplate(EHR_TEMPLATE_JINJA2_PATH)
    
    # Categorize genes into priority, standard, and unknown
    priority_genes = []
    standard_genes = []
    unknown_genes = []
    
    for g in pgx_genes:
        category = get_priority_category(g)
        gene_row = {
            "gene": g.get("gene", ""),
            "genotype": g.get("genotype", ""),
            "phenotype": g.get("metabolizer_status", "") or g.get("cpic_phenotype", ""),
            "medications": get_medication_examples(g.get("gene", ""))
        }
        
        if category == "priority":
            priority_genes.append(gene_row)
        elif category == "standard":
            standard_genes.append(gene_row)
        else:
            unknown_genes.append(gene_row)
    
    # Sort each list by gene name
    priority_genes.sort(key=lambda x: x["gene"])
    standard_genes.sort(key=lambda x: x["gene"])
    unknown_genes.sort(key=lambda x: x["gene"])
    
    logger.info(f"EHR Report (docxtpl): {len(priority_genes)} priority, {len(standard_genes)} standard, {len(unknown_genes)} unknown")
    
    # Build context for template rendering
    context = {
        # Patient info
        "patient_name": patient_info.get("patient_name", ""),
        "date_of_birth": patient_info.get("date_of_birth", ""),
        "mrn": patient_info.get("mrn", ""),
        "insurance": patient_info.get("insurance", ""),
        "report_date": patient_info.get("report_date", ""),
        "ordering_clinician": patient_info.get("ordering_clinician", ""),
        
        # Gene lists
        "priority_genes": priority_genes,
        "standard_genes": standard_genes,
        "unknown_genes": unknown_genes,
    }
    
    # Render the template with context
    template.render(context)
    
    # Save to file if path provided (for debugging)
    if output_path:
        template.save(output_path)
        logger.info(f"EHR report (docxtpl) saved to: {output_path}")
    
    # Return as bytes
    buffer = io.BytesIO()
    template.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
