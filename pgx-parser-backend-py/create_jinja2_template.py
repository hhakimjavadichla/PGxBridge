"""
Script to convert Note template.docx to a Jinja2-enabled template for docxtpl.

This script:
1. Reads the original template
2. Analyzes its structure (tables, paragraphs)
3. Creates a new template with Jinja2 placeholders
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches
from copy import deepcopy

# Paths
ORIGINAL_TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "pgx-parser-ui", "temp", "Note template.docx")
OUTPUT_TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "templates", "Note_template_jinja2.docx")

def analyze_template(doc_path):
    """Analyze the structure of the template."""
    doc = Document(doc_path)
    
    print("=" * 60)
    print("TEMPLATE ANALYSIS")
    print("=" * 60)
    
    # Analyze paragraphs
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  Para {i}: '{para.text[:80]}...' " if len(para.text) > 80 else f"  Para {i}: '{para.text}'")
    
    # Analyze tables
    print(f"\nTotal tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"\n  Table {i}: {len(table.rows)} rows x {len(table.columns)} columns")
        for j, row in enumerate(table.rows):
            cells = [cell.text[:30].replace('\n', ' ') for cell in row.cells]
            print(f"    Row {j}: {cells}")

def create_jinja2_template(input_path, output_path):
    """Create a Jinja2-enabled template from the original."""
    doc = Document(input_path)
    
    print("\n" + "=" * 60)
    print("CREATING JINJA2 TEMPLATE")
    print("=" * 60)
    
    # Table 0: Patient Info
    # Expected: Name, DOB, MRN, Insurance in first row
    if len(doc.tables) > 0:
        table0 = doc.tables[0]
        print(f"\nTable 0 (Patient Info): {len(table0.rows)} rows")
        if len(table0.rows) > 0:
            row = table0.rows[0]
            # Replace with Jinja2 placeholders
            if len(row.cells) >= 1:
                row.cells[0].text = "Name: {{patient_name}}"
            if len(row.cells) >= 2:
                row.cells[1].text = "DOB: {{date_of_birth}}"
            if len(row.cells) >= 3:
                row.cells[2].text = "MRN: {{mrn}}"
            if len(row.cells) >= 4:
                row.cells[3].text = "Insurance: {{insurance}}"
            print("  -> Added patient info placeholders")
    
    # Table 1: Priority Results (4 columns)
    # Keep header, replace data rows with loop
    if len(doc.tables) > 1:
        table1 = doc.tables[1]
        print(f"\nTable 1 (Priority Results): {len(table1.rows)} rows x {len(table1.columns)} cols")
        
        # Clear all rows except header
        while len(table1.rows) > 1:
            table1._tbl.remove(table1.rows[-1]._tr)
        
        # Add template row with Jinja2 loop
        row = table1.add_row()
        row.cells[0].text = "{%tr for gene in priority_genes %}"
        row.cells[1].text = ""
        row.cells[2].text = ""
        if len(row.cells) > 3:
            row.cells[3].text = ""
        
        # Add data row
        row = table1.add_row()
        row.cells[0].text = "{{gene.gene}}"
        row.cells[1].text = "{{gene.genotype}}"
        row.cells[2].text = "{{gene.phenotype}}"
        if len(row.cells) > 3:
            row.cells[3].text = "{{gene.medications}}"
        
        # Add end loop row
        row = table1.add_row()
        row.cells[0].text = "{%tr endfor %}"
        row.cells[1].text = ""
        row.cells[2].text = ""
        if len(row.cells) > 3:
            row.cells[3].text = ""
        
        print("  -> Added priority_genes loop")
    
    # Table 2: Standard Results (3 columns)
    if len(doc.tables) > 2:
        table2 = doc.tables[2]
        print(f"\nTable 2 (Standard Results): {len(table2.rows)} rows x {len(table2.columns)} cols")
        
        # Clear all rows except header
        while len(table2.rows) > 1:
            table2._tbl.remove(table2.rows[-1]._tr)
        
        # Add template row with Jinja2 loop
        row = table2.add_row()
        row.cells[0].text = "{%tr for gene in standard_genes %}"
        row.cells[1].text = ""
        row.cells[2].text = ""
        
        # Add data row
        row = table2.add_row()
        row.cells[0].text = "{{gene.gene}}"
        row.cells[1].text = "{{gene.genotype}}"
        row.cells[2].text = "{{gene.phenotype}}"
        
        # Add end loop row
        row = table2.add_row()
        row.cells[0].text = "{%tr endfor %}"
        row.cells[1].text = ""
        row.cells[2].text = ""
        
        print("  -> Added standard_genes loop")
    
    # Table 3: Unknown Results (3 columns)
    if len(doc.tables) > 3:
        table3 = doc.tables[3]
        print(f"\nTable 3 (Unknown Results): {len(table3.rows)} rows x {len(table3.columns)} cols")
        
        # Clear all rows except header
        while len(table3.rows) > 1:
            table3._tbl.remove(table3.rows[-1]._tr)
        
        # Add template row with Jinja2 loop
        row = table3.add_row()
        row.cells[0].text = "{%tr for gene in unknown_genes %}"
        row.cells[1].text = ""
        row.cells[2].text = ""
        
        # Add data row
        row = table3.add_row()
        row.cells[0].text = "{{gene.gene}}"
        row.cells[1].text = "{{gene.genotype}}"
        row.cells[2].text = "{{gene.phenotype}}"
        
        # Add end loop row
        row = table3.add_row()
        row.cells[0].text = "{%tr endfor %}"
        row.cells[1].text = ""
        row.cells[2].text = ""
        
        print("  -> Added unknown_genes loop")
    
    # Save the new template
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\n✓ Jinja2 template saved to: {output_path}")
    
    return output_path


if __name__ == "__main__":
    if not os.path.exists(ORIGINAL_TEMPLATE):
        print(f"ERROR: Template not found at {ORIGINAL_TEMPLATE}")
        sys.exit(1)
    
    # First analyze the template
    analyze_template(ORIGINAL_TEMPLATE)
    
    # Then create the Jinja2 version
    create_jinja2_template(ORIGINAL_TEMPLATE, OUTPUT_TEMPLATE)
    
    print("\n" + "=" * 60)
    print("NEXT STEPS:")
    print("=" * 60)
    print("1. Open the generated template in Word to verify formatting")
    print("2. The template uses {%tr for ... %} for table row loops")
    print("3. Variables: {{patient_name}}, {{date_of_birth}}, etc.")
    print("4. Gene loops: priority_genes, standard_genes, unknown_genes")
