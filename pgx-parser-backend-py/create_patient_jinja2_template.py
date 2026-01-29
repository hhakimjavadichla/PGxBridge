"""
Script to convert PGx Patient Report template to a Jinja2-enabled template for docxtpl.

Template structure:
- Table 0: Patient info (Name, DOB, Laboratory, Date of Test, Date of Interpretation)
- Table 1: Visual diagram (skip)
- Table 2: Intro text (skip)
- Table 3: High-risk/actionable results (4 columns: Gene, Genotype, Phenotype, Medications)
- Table 4: All PGx results (3 columns: Gene, Genotype, Phenotype)
- Tables 5-6: Footer/signature areas
"""

import os
import sys
from docx import Document

# Paths
ORIGINAL_TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "templates", "PGx Patient Report_04.08.25.docx")
OUTPUT_TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "templates", "PGx_Patient_Report_jinja2.docx")

def analyze_template(doc_path):
    """Analyze the structure of the template."""
    doc = Document(doc_path)
    
    print("=" * 60)
    print("PATIENT REPORT TEMPLATE ANALYSIS")
    print("=" * 60)
    
    # Analyze tables
    print(f"\nTotal tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"\n  Table {i}: {len(table.rows)} rows x {len(table.columns)} columns")
        for j, row in enumerate(table.rows[:3]):  # Show first 3 rows only
            cells = [cell.text[:30].replace('\n', ' ').strip() for cell in row.cells]
            print(f"    Row {j}: {cells}")
        if len(table.rows) > 3:
            print(f"    ... ({len(table.rows) - 3} more rows)")

def create_jinja2_template(input_path, output_path):
    """Create a Jinja2-enabled template from the original."""
    doc = Document(input_path)
    
    print("\n" + "=" * 60)
    print("CREATING JINJA2 PATIENT REPORT TEMPLATE")
    print("=" * 60)
    
    # === Table 0: Patient Info ===
    if len(doc.tables) > 0:
        table0 = doc.tables[0]
        print(f"\nTable 0 (Patient Info): {len(table0.rows)} rows")
        
        # Row 0: Patient Name (cell 1)
        if len(table0.rows) > 0 and len(table0.rows[0].cells) > 1:
            table0.rows[0].cells[1].text = "{{patient_name}}"
        
        # Row 1: DOB (cell 1), Date of Test (cell 3)
        if len(table0.rows) > 1:
            if len(table0.rows[1].cells) > 1:
                table0.rows[1].cells[1].text = "{{date_of_birth}}"
            if len(table0.rows[1].cells) > 3:
                table0.rows[1].cells[3].text = "{{report_date}}"
        
        # Row 2: Date of Interpretation (cell 3)
        if len(table0.rows) > 2 and len(table0.rows[2].cells) > 3:
            table0.rows[2].cells[3].text = "{{report_date}}"
        
        print("  -> Added patient info placeholders")
    
    # === Table 3: High-risk/Actionable Results (4 columns) ===
    if len(doc.tables) > 3:
        table3 = doc.tables[3]
        print(f"\nTable 3 (High-risk Results): {len(table3.rows)} rows x {len(table3.columns)} cols")
        
        # Clear all rows except header
        while len(table3.rows) > 1:
            table3._tbl.remove(table3.rows[-1]._tr)
        
        # Add loop start row
        row = table3.add_row()
        row.cells[0].text = "{%tr for gene in high_risk_genes %}"
        row.cells[1].text = ""
        row.cells[2].text = ""
        if len(row.cells) > 3:
            row.cells[3].text = ""
        
        # Add data row
        row = table3.add_row()
        row.cells[0].text = "{{gene.gene}}"
        row.cells[1].text = "{{gene.genotype}}"
        row.cells[2].text = "{{gene.phenotype}}"
        if len(row.cells) > 3:
            row.cells[3].text = "{{gene.medications}}"
        
        # Add loop end row
        row = table3.add_row()
        row.cells[0].text = "{%tr endfor %}"
        row.cells[1].text = ""
        row.cells[2].text = ""
        if len(row.cells) > 3:
            row.cells[3].text = ""
        
        print("  -> Added high_risk_genes loop")
    
    # === Table 4: All PGx Results (3 columns) ===
    if len(doc.tables) > 4:
        table4 = doc.tables[4]
        print(f"\nTable 4 (All Results): {len(table4.rows)} rows x {len(table4.columns)} cols")
        
        # Clear all rows except header
        while len(table4.rows) > 1:
            table4._tbl.remove(table4.rows[-1]._tr)
        
        # Add loop start row
        row = table4.add_row()
        row.cells[0].text = "{%tr for gene in all_genes %}"
        row.cells[1].text = ""
        row.cells[2].text = ""
        
        # Add data row
        row = table4.add_row()
        row.cells[0].text = "{{gene.gene}}"
        row.cells[1].text = "{{gene.genotype}}"
        row.cells[2].text = "{{gene.phenotype}}"
        
        # Add loop end row
        row = table4.add_row()
        row.cells[0].text = "{%tr endfor %}"
        row.cells[1].text = ""
        row.cells[2].text = ""
        
        print("  -> Added all_genes loop")
    
    # Save the new template
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\n✓ Jinja2 template saved to: {output_path}")


if __name__ == "__main__":
    if not os.path.exists(ORIGINAL_TEMPLATE):
        print(f"ERROR: Template not found at {ORIGINAL_TEMPLATE}")
        sys.exit(1)
    
    # Analyze the template
    analyze_template(ORIGINAL_TEMPLATE)
    
    # Create Jinja2 template
    create_jinja2_template(ORIGINAL_TEMPLATE, OUTPUT_TEMPLATE)
    
    print("\n" + "=" * 60)
    print("TEMPLATE VARIABLES:")
    print("=" * 60)
    print("Patient Info:")
    print("  - {{patient_name}}")
    print("  - {{date_of_birth}}")
    print("  - {{report_date}}")
    print("\nGene Lists:")
    print("  - high_risk_genes: Gene, Genotype, Phenotype, Medications")
    print("  - all_genes: Gene, Genotype, Phenotype")
